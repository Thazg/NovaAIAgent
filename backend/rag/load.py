import json
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

from config.settings import settings

BACKEND_DIR = Path(__file__).resolve().parents[1]
DATASET_DIR = Path(os.getenv("RAG_DATASET_DIR", BACKEND_DIR / "Dataset")).resolve()
UPLOADS_DIR = Path(settings.UPLOAD_FOLDER).resolve()
LEGACY_UPLOADS_DIR = (BACKEND_DIR / "uploads").resolve()

SUPPORTED_EXTENSIONS = {".pdf", ".md", ".markdown", ".rst", ".txt", ".py", ".ipynb", ".docx"}
TEXT_EXTENSIONS = {".md", ".markdown", ".rst", ".txt", ".py"}


@dataclass
class Document:
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)


def clean_text(text: str) -> str:
    if not text:
        return ""

    return "".join(
        ch for ch in text
        if not (0xD800 <= ord(ch) <= 0xDFFF)
    )


def file_metadata(file_path: Path) -> Dict[str, Any]:
    return {
        "source": str(file_path),
        "file_path": str(file_path),
        "file_name": file_path.name,
        "file_type": file_path.suffix.lower(),
    }


def add_metadata(text: str, file_path: Path) -> Document:
    return Document(text=clean_text(text), metadata=file_metadata(file_path))


def load_text(file_path: Path) -> List[Document]:
    text = file_path.read_text(encoding="utf-8", errors="ignore")
    return [add_metadata(text, file_path)]


def load_pdf(file_path: Path) -> List[Document]:
    if not PDF_AVAILABLE:
        return []
    reader = PdfReader(str(file_path))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text.strip())

    return [add_metadata("\n\n".join(pages), file_path)]


def load_notebook(file_path: Path) -> List[Document]:
    notebook = json.loads(file_path.read_text(encoding="utf-8", errors="ignore"))
    cells = []

    for cell in notebook.get("cells", []):
        source = cell.get("source", "")
        if isinstance(source, list):
            source = "".join(source)

        source = source.strip()
        if source:
            cells.append(source)

    return [add_metadata("\n\n".join(cells), file_path)]


def load_docx(file_path: Path) -> List[Document]:
    if not DOCX_AVAILABLE:
        return []
    document = DocxDocument(str(file_path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    return [add_metadata("\n\n".join(paragraphs), file_path)]


def load_file(file_path: Path) -> List[Document]:
    file_path = Path(file_path).resolve()
    suffix = file_path.suffix.lower()

    if suffix in TEXT_EXTENSIONS:
        documents = load_text(file_path)
    elif suffix == ".pdf":
        documents = load_pdf(file_path)
    elif suffix == ".ipynb":
        documents = load_notebook(file_path)
    elif suffix == ".docx":
        documents = load_docx(file_path)
    else:
        return []

    return documents


def iter_supported_files(input_dir=DATASET_DIR) -> List[Path]:
    input_dir = Path(input_dir).resolve()
    if not input_dir.exists():
        return []

    return sorted(
        file_path
        for file_path in input_dir.rglob("*")
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def load_documents(input_dir=DATASET_DIR) -> List[Document]:
    documents: List[Document] = []
    seen_paths: set[str] = set()

    for file_path in iter_supported_files(input_dir):
        resolved = str(file_path.resolve())
        if resolved in seen_paths:
            continue
        seen_paths.add(resolved)
        documents.extend(load_file(file_path))

    extra_dirs = []
    if UPLOADS_DIR.exists():
        extra_dirs.append(UPLOADS_DIR)
    if LEGACY_UPLOADS_DIR.exists() and LEGACY_UPLOADS_DIR not in extra_dirs:
        extra_dirs.append(LEGACY_UPLOADS_DIR)

    for extra_dir in extra_dirs:
        try:
            inside_dataset = extra_dir.resolve().is_relative_to(DATASET_DIR.resolve())
        except AttributeError:
            inside_dataset = str(extra_dir.resolve()).startswith(str(DATASET_DIR.resolve()))

        if inside_dataset:
            continue

        for file_path in iter_supported_files(extra_dir):
            resolved = str(file_path.resolve())
            if resolved in seen_paths:
                continue
            seen_paths.add(resolved)
            documents.extend(load_file(file_path))

    return documents


def print_load_summary(documents: List[Document], input_dir=DATASET_DIR) -> None:
    counts = {}

    for document in documents:
        file_type = document.metadata.get("file_type", "unknown")
        counts[file_type] = counts.get(file_type, 0) + 1

    print(f"Loaded {len(documents)} documents from: {Path(input_dir).resolve()}")
    for file_type, count in sorted(counts.items()):
        print(f"- {file_type}: {count}")


if __name__ == "__main__":
    documents = load_documents()
    print_load_summary(documents)
